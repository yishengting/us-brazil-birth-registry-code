#!/usr/bin/env python3
"""Apply reproducible A4 table geometry to BJOG submission table DOCX files.

Pandoc's default Word export gives every column nearly equal width.  That makes
the wider supplementary tables technically editable but visually unusable.
This post-processor keeps the underlying table text unchanged while applying
landscape geometry, repeat headers, non-splitting rows, and content-aware
column widths.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MAIN_READY = ROOT / "submission" / "tables" / "main" / "publication_ready"
SUPP_READY = ROOT / "submission" / "tables" / "supplementary" / "publication_ready"


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    for old_height in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(old_height)
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def header_width(header: str) -> float:
    h = re.sub(r"\s+", " ", header.strip().lower())
    if any(token in h for token in ("note", "rationale", "handling", "coding detail", "final decision")):
        return 2.35
    if any(token in h for token in ("source variables", "rule used", "definition", "age subtype", "risk profile")):
        return 1.55
    if any(token in h for token in ("analysis", "step", "domain", "characteristic")):
        return 1.35
    if "country" in h:
        return 1.0
    if "outcome" in h:
        return 1.15
    if "birth year" in h or h == "year":
        return 0.85
    if any(token in h for token in ("p value", "p for", "percent", "difference", "sample n", "model n", "births", "cells")):
        return 1.0
    if "a rr" in h or "arr" in h or "95% ci" in h or "risk score" in h:
        return 1.15
    return 1.15


def column_widths(table, usable_inches: float) -> list[float]:
    headers = [cell.text for cell in table.rows[0].cells]
    weights = [header_width(h) for h in headers]
    # Give genuinely narrative columns room even when their header is terse.
    for idx in range(len(headers)):
        values = [re.sub(r"\s+", " ", row.cells[idx].text.strip()) for row in table.rows[1:]]
        longest_word = max((len(word) for value in values for word in value.split()), default=0)
        longest_value = max((len(value) for value in values), default=0)
        if longest_value > 85:
            weights[idx] = max(weights[idx], 2.1)
        elif longest_value > 45:
            weights[idx] = max(weights[idx], 1.55)
        if longest_word > 18:
            weights[idx] = max(weights[idx], 1.35)

    total = sum(weights)
    widths = [usable_inches * weight / total for weight in weights]
    # Avoid columns so narrow that Word breaks ordinary words letter by letter.
    floor = 0.72 if len(widths) >= 8 else 0.85
    for _ in range(3):
        shortfall = sum(max(0.0, floor - width) for width in widths)
        widths = [max(floor, width) for width in widths]
        donors = [i for i, width in enumerate(widths) if width > floor + 0.25]
        if not donors or shortfall <= 0:
            break
        donor_room = sum(widths[i] - floor for i in donors)
        for i in donors:
            widths[i] -= shortfall * (widths[i] - floor) / donor_room
    scale = usable_inches / sum(widths)
    return [width * scale for width in widths]


def set_table_geometry(table, widths_inches: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total_dxa = round(sum(widths_inches) * 1440)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        if row_index == 0:
            set_repeat_header(row)
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, round(widths_inches[col_index] * 1440))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.3)
                    if row_index == 0:
                        run.bold = True


def format_supplement(path: Path):
    doc = Document(path)
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.keep_together = True
        if paragraph.style.name.startswith("Heading") or paragraph.style.name == "Title":
            paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            run.font.name = "Times New Roman"

    usable_inches = 11.69 - 1.2
    for table in doc.tables:
        set_table_geometry(table, column_widths(table, usable_inches))

    doc.save(path)
    print(f"Formatted {path}")


def format_main(path: Path):
    """Set an explicit A4 portrait page while preserving the audited main-table styling."""

    doc = Document(path)
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
    doc.save(path)
    print(f"Formatted {path}")


def main():
    main_targets = [
        MAIN_READY / "main_tables_publication_ready.docx",
        MAIN_READY / "table1_publication_ready.docx",
    ]
    targets = [SUPP_READY / "supplementary_tables_publication_ready.docx"]
    targets.extend(sorted(SUPP_READY.glob("Supplementary_Table_*_publication_ready.docx")))
    missing = [path for path in main_targets + targets if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing expected table DOCX: " + ", ".join(map(str, missing)))
    for path in main_targets:
        format_main(path)
    for path in targets:
        format_supplement(path)


if __name__ == "__main__":
    main()
